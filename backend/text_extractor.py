"""
Text Extraction Module
Handles PDF, DOCX, and TXT file processing
"""
from docx import Document   # This is correct after installing python-docx
import PyPDF2
import pdfplumber
from docx import Document
import io
from typing import Dict, Any

class TextExtractor:
    """
    Extracts text from various document formats
    Supports: PDF, DOCX, TXT
    """
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from PDF using PyPDF2 with fallback to pdfplumber
        
        Args:
            file_bytes: Raw PDF file bytes
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Primary extraction with PyPDF2 (faster)
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            total_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # If extraction fails, try pdfplumber (better for complex layouts)
            if len(text.strip()) < 100:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
            
            return {
                "text": text,
                "pages": total_pages,
                "extraction_method": "PyPDF2/pdfplumber",
                "success": True
            }
            
        except Exception as e:
            return {
                "text": "",
                "pages": 0,
                "extraction_method": "Failed",
                "success": False,
                "error": str(e)
            }
    
    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from DOCX files
        
        Args:
            file_bytes: Raw DOCX file bytes
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            
            # Extract tables if any
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    tables_text.append(row_text)
            
            if tables_text:
                text += "\n\nTables:\n" + "\n".join(tables_text)
            
            return {
                "text": text,
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "extraction_method": "python-docx",
                "success": True
            }
            
        except Exception as e:
            return {
                "text": "",
                "paragraphs": 0,
                "extraction_method": "Failed",
                "success": False,
                "error": str(e)
            }
    
    @staticmethod
    def extract_from_txt(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from TXT files
        
        Args:
            file_bytes: Raw TXT file bytes
            
        Returns:
            Dictionary containing extracted text
        """
        try:
            text = file_bytes.decode('utf-8')
            
            return {
                "text": text,
                "extraction_method": "UTF-8 decode",
                "success": True
            }
            
        except UnicodeDecodeError:
            # Try alternative encoding
            try:
                text = file_bytes.decode('latin-1')
                return {
                    "text": text,
                    "extraction_method": "Latin-1 decode",
                    "success": True
                }
            except Exception as e:
                return {
                    "text": "",
                    "extraction_method": "Failed",
                    "success": False,
                    "error": str(e)
                }
